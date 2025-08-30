# app.py
import os
import re
import sys
from datetime import datetime
import traceback

# Groq client
from groq import Groq

# Word & PDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import ParagraphStyle

# Gradio UI
import gradio as gr

# Initialize Groq client (expect GROQ_API_KEY in env)
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
if not GROQ_API_KEY:
    # we'll still allow UI to run but model calls will return an instructive error
    client = None
else:
    client = Groq(api_key=GROQ_API_KEY)

# In-memory store
book_data = {
    "title": "",
    "grade": "",
    "author": "",
    "developer": "Najaf Ali Sharqi",
    "medium": "English",
    "toc": [],         # list of lines like "1. Chapter title"
    "chapters": {},    # int -> chapter text
    "preface": ""
}


def clean_formatting(text: str) -> str:
    """Remove markdown, emojis and collapse whitespace."""
    if not text:
        return ""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    # remove common symbols/emojis used before
    text = re.sub(r"[#•●→✅🔹🔸📘📖📑📤📥⬇️🎉>]", "", text)
    # normalize whitespace
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_model_call(prompt: str, system: str = None, model: str = "llama3-70b-8192", temperature: float = 0.7):
    """Call Groq model safely and return string or raise informative error."""
    if client is None:
        raise RuntimeError("GROQ_API_KEY not found in environment. Set GROQ_API_KEY before calling the model.")
    try:
        messages = []
        if system:
            messages.append({"role": "system", "content": system})
        messages.append({"role": "user", "content": prompt})
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature
        )
        return response.choices[0].message.content
    except Exception as e:
        tb = traceback.format_exc()
        raise RuntimeError(f"Model call failed: {e}\n{tb}")


def parse_preface_and_toc(raw_output: str):
    """Try to split model output into preface and Table of Contents robustly."""
    if not raw_output:
        return "", []
    # Search for a line that starts "Table of Contents" case-insensitively
    m = re.search(r"Table of Contents", raw_output, re.IGNORECASE)
    if m:
        preface_part = raw_output[:m.start()].strip()
        toc_part = raw_output[m.end():].strip()
    else:
        # fallback: attempt to heuristically split on 'Chapter 1' or similar
        m2 = re.search(r"(?mi)^(Chapter\s*1|1\.)", raw_output)
        if m2:
            preface_part = raw_output[:m2.start()].strip()
            toc_part = raw_output[m2.start():].strip()
        else:
            # No clear TOC — small fallback
            preface_part = raw_output.strip()
            toc_part = ""

    # Build toc list
    toc_lines = []
    if toc_part:
        for line in toc_part.splitlines():
            line = line.strip()
            if not line:
                continue
            # If the model included "Table of Contents" header in the remaining text, skip it
            if re.match(r"^\s*Table of Contents\s*$", line, re.IGNORECASE):
                continue
            toc_lines.append(line)
    else:
        # generate conservative default 7-chapter TOC placeholders
        toc_lines = [f"{i}. Chapter {i}" for i in range(1, 8)]

    return clean_formatting(preface_part), [clean_formatting(l) for l in toc_lines]


def generate_book_intro(title: str, grade: str, author: str, medium: str):
    """Generate preface and initial TOC using model."""
    title = (title or "").strip()
    grade = (grade or "").strip()
    author = (author or "").strip()
    medium = medium or "English"

    book_data.update({"title": title, "grade": grade, "author": author, "medium": medium})

    prompt = f"""
You are an expert Pakistani educational textbook author. Language/medium: {medium}.
Write a formal Preface (APA 7th edition tone) and a clear Table of Contents for a beginner-level textbook titled "{title}" for Grade {grade}, authored by {author}.
Requirements:
- Preface: concise, formal, APA tone (do NOT include markdown or emojis).
- Table of Contents: exactly 7 chapter titles, each short and appropriate for absolute beginners in Pakistan.
Return the Preface followed by a heading 'Table of Contents' and then the 7 lines of TOC.
Do NOT produce chapter bodies.
"""
    try:
        raw_out = safe_model_call(prompt)
        preface, toc = parse_preface_and_toc(raw_out)
        book_data["preface"] = preface
        book_data["toc"] = toc
        # clear any previously generated chapters
        book_data["chapters"] = {}
        return f"Preface:\n\n{preface}\n\nTable of Contents:\n\n" + "\n".join(toc)
    except Exception as e:
        return f"Error generating book intro: {e}"


def generate_chapter(ch_num):
    """Generate one chapter (with SLOs, aligned content, examples, glossary, 10 MCQs)."""
    try:
        ch_index = int(ch_num)
    except Exception:
        return "Invalid chapter number."

    if ch_index in book_data["chapters"]:
        return book_data["chapters"][ch_index]

    # obtain chapter title from toc if present
    ch_title = f"Chapter {ch_index}"
    if 1 <= ch_index <= len(book_data.get("toc", [])):
        # strip leading numbering if present
        raw = book_data["toc"][ch_index - 1]
        # remove '1.' or 'Chapter 1 -' prefixes
        raw = re.sub(r"^\s*\d+[\.\-\:]\s*", "", raw)
        raw = re.sub(r"(?i)^chapter\s*\d+[\.\-\:]*\s*", "", raw)
        ch_title = raw.strip() or ch_title

    prompt = f"""
You are an expert Pakistani textbook writer. Medium: {book_data['medium']}.
Write Chapter {ch_index}: "{ch_title}" for the textbook titled "{book_data['title']}" (Grade {book_data['grade']}) by {book_data['author']}.
Structure the chapter exactly with these sections (in this order):

1. Student Learning Outcomes (5 to 7 short SLO statements).
2. For each SLO, provide detailed content aligned to that SLO. For every concept include:
   - What (definition in one or two sentences)
   - Why (why it matters; one or two sentences)
   - How (a simple step/example or classroom activity) — use at least one Pakistan-relevant example.
3. Activities or short case studies (1–2) with instructions for students and teacher prompts.
4. Glossary (key terms and brief definitions).
5. Post-assessment: 10 MCQs (without answers). Make them beginner-level and relevant to Pakistan context.

Styling notes: no markdown, no emojis. Use clear simple language for absolute beginners. Keep each section labeled (e.g., "Student Learning Outcomes:").
Return the whole chapter as plain text.
"""
    try:
        raw_out = safe_model_call(prompt)
        chapter_text = clean_formatting(raw_out)
        book_data["chapters"][ch_index] = chapter_text
        return chapter_text
    except Exception as e:
        return f"Error generating chapter: {e}"


def generate_all_chapters():
    """Generate all chapters according to current TOC length (default 7)."""
    results = []
    total = max(1, len(book_data.get("toc") or []))
    for i in range(1, total + 1):
        results.append(f"Generating Chapter {i}...")
        chapter_text = generate_chapter(i)
        results.append(chapter_text[:1000] + ("..." if len(chapter_text) > 1000 else ""))
    return "\n\n".join(results)


# --- Word export helpers ---
def _add_centered_paragraph(doc: Document, text: str, size: int = 12, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    try:
        run.font.name = "Times New Roman"
    except Exception:
        pass
    p.paragraph_format.line_spacing = 1.5
    return p


def _add_left_paragraph(doc: Document, text: str, size: int = 12, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    try:
        run.font.name = "Times New Roman"
    except Exception:
        pass
    p.paragraph_format.line_spacing = 1.5
    return p


def _insert_word_toc(doc: Document):
    """
    Insert a Word Table of Contents field. Word will populate page numbers when user updates fields (References -> Update Table).
    This is the practical way to provide a final TOC with page numbers for .docx files.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # TOC field to include levels 1-3, use hyperlinks, hide tab leader in web layout
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    # A "please update" placeholder that Word will replace
    run2 = paragraph.add_run("Right-click the table and choose 'Update Field' to populate page numbers.")
    run2.font.size = Pt(10)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)


def export_book_word():
    """Create a .docx file and return path."""
    try:
        doc = Document()
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(12)

        # Developer at top center
        _add_centered_paragraph(doc, book_data.get("developer", ""), size=12, bold=True)

        # Title page
        _add_centered_paragraph(doc, book_data.get("title", ""), size=18, bold=True)
        _add_centered_paragraph(doc, f"Grade: {book_data.get('grade','')}", size=12)
        _add_centered_paragraph(doc, f"Author: {book_data.get('author','')}", size=12)
        _add_centered_paragraph(doc, f"Developed by: {book_data.get('developer','')}", size=12)
        _add_centered_paragraph(doc, f"Date: {datetime.now().strftime('%B %d, %Y')}", size=12)
        doc.add_page_break()

        # Preface
        _add_centered_paragraph(doc, "Preface", size=12, bold=True)
        _add_left_paragraph(doc, book_data.get("preface", ""))
        doc.add_page_break()

        # Table of Contents (field)
        _add_centered_paragraph(doc, "Table of Contents", size=12, bold=True)
        _insert_word_toc(doc)
        doc.add_page_break()

        # Chapters
        for num in sorted(book_data.get("chapters", {}).keys()):
            # chapter title: take from toc if available
            title_line = book_data["toc"][num - 1] if 0 <= num - 1 < len(book_data.get("toc", [])) else f"Chapter {num}"
            # chapter heading as 12 bold center (per your spec for subheadings)
            _add_centered_paragraph(doc, title_line, size=12, bold=True)
            # write the chapter lines
            for line in book_data["chapters"][num].splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                # mark section headers (Student Learning Outcomes, Activities, Glossary, Assessment) as bold left
                if re.match(r"(?i)^(student learning outcomes|activities|assessment|glossary|post-assessment|activities or case studies)\b", stripped):
                    _add_left_paragraph(doc, stripped, size=12, bold=True)
                else:
                    _add_left_paragraph(doc, stripped, size=12, bold=False)
            doc.add_page_break()

        out_path = "/tmp/Textbook_AI_Generated.docx"
        doc.save(out_path)
        return out_path
    except Exception as e:
        return f"Error exporting Word: {e}\n{traceback.format_exc()}"


def export_book_pdf():
    """Create a PDF (simpler TOC — Word is recommended for final page-numbered TOC)."""
    try:
        out_path = "/tmp/Textbook_AI_Generated.pdf"
        doc = SimpleDocTemplate(out_path, pagesize=A4)
        normal = ParagraphStyle("Normal", fontName="Times-Roman", fontSize=12, leading=18)
        heading_center = ParagraphStyle("HCenter", fontName="Times-Bold", fontSize=18, leading=22, alignment=1)
        subheading = ParagraphStyle("Sub", fontName="Times-Bold", fontSize=12, leading=16)

        story = []
        # Developer and title
        story.append(Paragraph(book_data.get("developer", ""), subheading))
        story.append(Paragraph(book_data.get("title", ""), heading_center))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Grade: {book_data.get('grade','')}", normal))
        story.append(Paragraph(f"Author: {book_data.get('author','')}", normal))
        story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", normal))
        story.append(PageBreak())

        # Preface
        story.append(Paragraph("Preface", subheading))
        story.append(Paragraph(book_data.get("preface", ""), normal))
        story.append(PageBreak())

        # TOC (no automatic page numbers in this PDF)
        story.append(Paragraph("Table of Contents (open the Word .docx to get page-numbered TOC)", subheading))
        for line in book_data.get("toc", []):
            story.append(Paragraph(line, normal))
        story.append(PageBreak())

        # Chapters
        for num in sorted(book_data.get("chapters", {}).keys()):
            title_line = book_data["toc"][num - 1] if 0 <= num - 1 < len(book_data.get("toc", [])) else f"Chapter {num}"
            story.append(Paragraph(title_line, subheading))
            for line in book_data["chapters"][num].splitlines():
                if not line.strip():
                    continue
                story.append(Paragraph(line.strip(), normal))
            story.append(PageBreak())

        doc.build(story)
        return out_path
    except Exception as e:
        return f"Error exporting PDF: {e}\n{traceback.format_exc()}"


# --- Gradio UI ---

def in_colab():
    try:
        import google.colab
        return True
    except Exception:
        return False


with gr.Blocks(title="AI Textbook Generator - APA Style") as demo:
    # Top header with developer name centered
    header_html = f"""
    <div style="text-align:center; margin-bottom:6px;">
        <h2 style="margin:0;">📘 AI Textbook Generator – APA Style</h2>
        <h4 style="margin:0;">Developed by: {book_data['developer']}</h4>
    </div>
    """
    gr.HTML(header_html)

    with gr.Row():
        with gr.Column(scale=2):
            title_in = gr.Textbox(label="Book Title", placeholder="e.g., Introduction to Environmental Science", lines=1)
            grade_in = gr.Textbox(label="Grade Level", placeholder="e.g., 9", lines=1)
            author_in = gr.Textbox(label="Author Name", placeholder="e.g., Fatima Khan", lines=1)
            medium_in = gr.Radio(choices=["English", "Urdu"], value="English", label="Medium")

            gen_intro_btn = gr.Button("🧠 Generate TOC + Preface")
            intro_out = gr.Textbox(label="Generated Preface + TOC", lines=12)

            gen_all_btn = gr.Button("Generate All Chapters (1..TOC)")
            gen_all_out = gr.Textbox(label="Generate All Chapters Output (log)", lines=10)

            # Chapter generation controls
            gr.Markdown("### Generate single chapter")
            chapter_out = gr.Textbox(label="Chapter output", lines=20)
            # create 7 pairs of (button, hidden state) dynamically
            chapter_buttons = []
            chapter_states = []
            for i in range(1, 8):
                st = gr.State(i)
                btn = gr.Button(f"Generate Chapter {i}")
                btn.click(fn=generate_chapter, inputs=[st], outputs=chapter_out)
                chapter_buttons.append(btn)
                chapter_states.append(st)

        with gr.Column(scale=1):
            gr.Markdown("#### Actions & Download")
            gen_all_btn.click(fn=generate_all_chapters, inputs=[], outputs=gen_all_out)
            word_btn = gr.Button("Download MS Word (.docx)")
            pdf_btn = gr.Button("Download PDF (.pdf)")
            file_out = gr.File()
            word_btn.click(fn=export_book_word, inputs=None, outputs=file_out)
            pdf_btn.click(fn=export_book_pdf, inputs=None, outputs=file_out)

            gr.Markdown("#### Notes")
            gr.Markdown("- Make sure `GROQ_API_KEY` environment variable is set to call the model.")
            gr.Markdown("- For final page-numbered Table of Contents: open the downloaded `.docx` in Microsoft Word and update the TOC (References → Update Table). Word will populate page numbers.")
            gr.Markdown("- If testing on Google Colab: the app will run with `share=True` automatically.")

    # Wire generate intro button
    gen_intro_btn.click(fn=generate_book_intro, inputs=[title_in, grade_in, author_in, medium_in], outputs=intro_out)

# Launch logic
if __name__ == "__main__":
    port_env = int(os.environ.get("PORT", 7860))
    if in_colab():
        # In Colab we prefer share=True
        demo.launch(share=True)
    else:
        # On Railway / production bind to 0.0.0.0 and given port
        demo.launch(server_name="0.0.0.0", server_port=port_env, share=False)
